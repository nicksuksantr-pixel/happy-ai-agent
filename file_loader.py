"""
file_loader.py — Read all file types user uploads
Supports: image, PDF, Word, Excel, CSV, text
"""
import io
from pathlib import Path
from typing import List, Dict, Tuple

# v2.5.1 (Cos audit B-03): hoisted from inside build_gemini_parts() so an
# ImportError surfaces at module load (where it's diagnosable) instead of
# deferring until the first multimodal call deep in the pipeline.
from google.genai import types as _genai_types

IMAGE_TYPES = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
PDF_TYPES = {".pdf"}
WORD_TYPES = {".docx"}
EXCEL_TYPES = {".xlsx", ".xls"}
CSV_TYPES = {".csv"}
TEXT_TYPES = {".txt", ".md"}
ALL_SUPPORTED = IMAGE_TYPES | PDF_TYPES | WORD_TYPES | EXCEL_TYPES | CSV_TYPES | TEXT_TYPES


def get_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext in IMAGE_TYPES: return "image"
    if ext in PDF_TYPES:   return "pdf"
    if ext in WORD_TYPES:  return "word"
    if ext in EXCEL_TYPES: return "excel"
    if ext in CSV_TYPES:   return "csv"
    if ext in TEXT_TYPES:  return "text"
    return "unknown"


def is_supported(filename: str) -> bool:
    return get_file_type(filename) != "unknown"


def get_mime_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".csv": "text/csv",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }
    return mime_map.get(ext, "application/octet-stream")


def extract_text_from_word(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        return "[ERROR] python-docx not installed"
    except Exception as e:
        return f"[ERROR reading Word: {e}]"


def extract_text_from_excel(file_bytes: bytes, max_rows: int = 50) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"### Sheet: {sheet_name}")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count >= max_rows:
                    parts.append(f"... (truncated at {max_rows} rows)")
                    break
                row_text = " | ".join(str(c) if c is not None else "" for c in row)
                if row_text.strip(" |"):
                    parts.append(row_text)
                    row_count += 1
            parts.append("")
        wb.close()
        return "\n".join(parts)
    except ImportError:
        return "[ERROR] openpyxl not installed"
    except Exception as e:
        return f"[ERROR reading Excel: {e}]"


def extract_text_from_csv(file_bytes: bytes, max_rows: int = 50) -> str:
    try:
        import csv
        text = file_bytes.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(text))
        parts = []
        for i, row in enumerate(reader):
            if i >= max_rows:
                parts.append(f"... (truncated at {max_rows} rows)")
                break
            parts.append(" | ".join(row))
        return "\n".join(parts)
    except Exception as e:
        return f"[ERROR reading CSV: {e}]"


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"### Page {i+1}")
                parts.append(text)
        return "\n\n".join(parts)
    except ImportError:
        return "[ERROR] pypdf not installed"
    except Exception as e:
        return f"[ERROR reading PDF: {e}]"


def load_file_for_gemini(filename: str, file_bytes: bytes) -> Dict:
    file_type = get_file_type(filename)
    size = len(file_bytes)
    
    if file_type == "image":
        return {
            "filename": filename, "type": "image", "size": size,
            "send_as": "inline_data", "mime_type": get_mime_type(filename),
            "data": file_bytes,
            "preview": f"[Image: {filename}, {size:,} bytes]",
        }
    elif file_type == "pdf":
        if size <= 20 * 1024 * 1024:
            return {
                "filename": filename, "type": "pdf", "size": size,
                "send_as": "inline_data", "mime_type": "application/pdf",
                "data": file_bytes,
                "preview": f"[PDF: {filename}, {size:,} bytes]",
            }
        else:
            text = extract_text_from_pdf(file_bytes)
            return {
                "filename": filename, "type": "pdf", "size": size,
                "send_as": "text",
                "data": f"# Content from {filename} (large PDF, text extracted)\n\n{text}",
                "preview": f"[PDF: {filename}, {size:,} bytes - extracted as text]",
            }
    elif file_type == "word":
        text = extract_text_from_word(file_bytes)
        return {
            "filename": filename, "type": "word", "size": size, "send_as": "text",
            "data": f"# Content from {filename} (Word document)\n\n{text}",
            "preview": f"[Word: {filename}, {size:,} bytes]",
        }
    elif file_type == "excel":
        text = extract_text_from_excel(file_bytes)
        return {
            "filename": filename, "type": "excel", "size": size, "send_as": "text",
            "data": f"# Content from {filename} (Excel spreadsheet)\n\n{text}",
            "preview": f"[Excel: {filename}, {size:,} bytes]",
        }
    elif file_type == "csv":
        text = extract_text_from_csv(file_bytes)
        return {
            "filename": filename, "type": "csv", "size": size, "send_as": "text",
            "data": f"# Content from {filename} (CSV data)\n\n{text}",
            "preview": f"[CSV: {filename}, {size:,} bytes]",
        }
    elif file_type == "text":
        text = file_bytes.decode("utf-8", errors="ignore")
        return {
            "filename": filename, "type": "text", "size": size, "send_as": "text",
            "data": f"# Content from {filename}\n\n{text}",
            "preview": f"[Text: {filename}, {size:,} bytes]",
        }
    else:
        return {
            "filename": filename, "type": "unknown", "size": size, "send_as": "text",
            "data": f"[Unsupported file type: {filename}]",
            "preview": f"[Unknown: {filename}]",
        }


def build_gemini_parts(text_prompt: str, files: List[Dict]) -> list:
    parts = [text_prompt]
    for f in files:
        if f["send_as"] == "inline_data":
            parts.append(_genai_types.Part.from_bytes(
                data=f["data"], mime_type=f["mime_type"],
            ))
        else:
            parts.append(f"\n\n{f['data']}")
    return parts


def save_attachments_to_session(session_path: Path, files: List[Tuple[str, bytes]]) -> List[str]:
    if not files:
        return []
    attach_dir = session_path / "attachments"
    attach_dir.mkdir(exist_ok=True)
    saved = []
    for filename, data in files:
        safe_name = Path(filename).name
        target = attach_dir / safe_name
        counter = 1
        while target.exists():
            stem = Path(safe_name).stem
            suffix = Path(safe_name).suffix
            target = attach_dir / f"{stem}_{counter}{suffix}"
            counter += 1
        target.write_bytes(data)
        saved.append(str(target.relative_to(session_path)))
    return saved


def load_attachments_from_session(session_path: Path) -> List[Dict]:
    # v2.8.0 (Cos audit B-13): the bare `except Exception: continue`
    # silently swallowed every failure — corrupt files, format-version
    # mismatches, unknown bugs. Surface them via print so they show up
    # in HAPPY's crash.log tee + don't hide the root cause. Caller can
    # still iterate the survivors.
    attach_dir = session_path / "attachments"
    if not attach_dir.exists():
        return []
    files = []
    for f in sorted(attach_dir.iterdir()):
        if not f.is_file():
            continue
        try:
            data = f.read_bytes()
            files.append(load_file_for_gemini(f.name, data))
        except Exception as e:
            try:
                print(
                    f"[file_loader] skipping unreadable attachment "
                    f"{f.name}: {type(e).__name__}: {str(e)[:200]}",
                    flush=True,
                )
            except Exception:
                pass
            continue
    return files